from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

doc = Document()

# ---- Page setup (RTL) ----
for section in doc.sections:
    section.right_to_left = True
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.right_margin = Cm(2.5)
    section.left_margin = Cm(2.5)

# ---- Default style ----
style = doc.styles['Normal']
style.font.name = 'Arial'
style.font.size = Pt(11)
style.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.3
rPr = style.element.get_or_add_rPr()
rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:cs="Arial"/>')
rPr.append(rFonts)

# ---- Colors ----
BLACK = RGBColor(0x00, 0x00, 0x00)
DARK = RGBColor(0x1A, 0x1A, 0x1A)
DARK_BLUE = RGBColor(0x0D, 0x47, 0xA1)
DARK_GREEN = RGBColor(0x1B, 0x7A, 0x3D)
DARK_RED = RGBColor(0xB7, 0x1C, 0x1C)
DARK_ORANGE = RGBColor(0xBF, 0x6A, 0x00)
GRAY = RGBColor(0x44, 0x44, 0x44)
ACCENT = RGBColor(0x0D, 0x47, 0xA1)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
TABLE_HEADER_BG = '1A344D'
TABLE_ALT_BG = 'E8F0FE'

# ---- Helpers ----
def add_heading_styled(text, level=1):
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in h.runs:
        run.font.color.rgb = DARK_BLUE
        if level == 1:
            run.font.size = Pt(20)
        elif level == 2:
            run.font.size = Pt(15)
        elif level == 3:
            run.font.size = Pt(12)
    return h

def add_para(text, bold=False, color=DARK, size=None, align_rtl=True):
    p = doc.add_paragraph()
    if align_rtl:
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(text)
    if bold:
        run.bold = True
    run.font.color.rgb = color
    if size:
        run.font.size = Pt(size)
    return p

def add_bullet(text, level=0, color=DARK, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.right_indent = Cm(1 + level * 0.5)
    if bold_prefix:
        run_b = p.add_run(bold_prefix)
        run_b.bold = True
        run_b.font.color.rgb = color
        run = p.add_run(text)
        run.font.color.rgb = color
    else:
        p.clear()
        run = p.add_run(text)
        run.font.color.rgb = color
    return p

def add_numbered(text, number, color=DARK, bold=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.right_indent = Cm(1)
    run = p.add_run(f'{number}. ')
    run.bold = True
    run.font.color.rgb = DARK_BLUE
    run2 = p.add_run(text)
    run2.font.color.rgb = color
    if bold:
        run2.bold = True
    return p

def set_cell_shading(cell, color):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def add_cell_text(cell, text, bold=False, color=DARK, size=9, align_rtl=True):
    cell.text = ''
    p = cell.paragraphs[0]
    if align_rtl:
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    if bold:
        run.bold = True

def create_table(headers, data, col_widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.RIGHT
    table.style = 'Table Grid'

    # Header row
    for i, txt in enumerate(headers):
        add_cell_text(table.rows[0].cells[i], txt, bold=True, color=WHITE, size=9)
        set_cell_shading(table.rows[0].cells[i], TABLE_HEADER_BG)

    # Data rows
    for row_idx, row_data in enumerate(data):
        row = table.add_row()
        bg = TABLE_ALT_BG if row_idx % 2 == 0 else 'FFFFFF'
        for i, txt in enumerate(row_data):
            add_cell_text(row.cells[i], txt, color=DARK, size=9)
            set_cell_shading(row.cells[i], bg)

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    return table

# ================================================================
# DOCUMENT CONTENT
# ================================================================

# ---- Cover Page ----
for _ in range(6):
    doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('خطة تطوير واجهة المستخدم\nPhase 1B')
run.font.size = Pt(30)
run.bold = True
run.font.color.rgb = DARK_BLUE

doc.add_paragraph()
subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('توحيد Sidebar - Header - FloatingBell - AppShell\n')
run.font.size = Pt(16)
run.font.color.rgb = DARK

run2 = subtitle.add_run('نمط التطوير التدريجي: Strangler Fig Pattern')
run2.font.size = Pt(13)
run2.font.color.rgb = GRAY
run2.italic = True

doc.add_paragraph()
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = meta.add_run('منصة CyberSecurty v11\nتاريخ الخطة: 19 يونيو 2026')
run.font.size = Pt(12)
run.font.color.rgb = GRAY

doc.add_page_break()

# ---- TOC ----
add_heading_styled('فهرس المحتويات', level=1)
toc = [
    '1. نظرة عامة وأهداف الخطة',
    '2. استراتيجية Strangler Fig Pattern',
    '3. المرحلة الأولى — التحضير والتأهب (Preparation)',
    '4. المرحلة الثانية — بناء UnifiedSidebar (جنبًا إلى جنب)',
    '5. المرحلة الثالثة — بناء useNotifications Hook',
    '6. المرحلة الرابعة — بناء BellNotifications Component',
    '7. المرحلة الخامسة — بناء UnifiedHeader',
    '8. المرحلة السادسة — بناء AppShell',
    '9. المرحلة السابعة — الهجرة التدريجية (صفحة صفحة)',
    '10. المرحلة الثامنة — حذف الأنظمة القديمة',
    '11. المرحلة التاسعة — اختبار شامل وتدقيق أمني',
    '12. خطة الطوارئ والاسترجاع (Rollback Plan)',
    '13. الجدول الزمني المقدر',
    '14. الملخص',
]
for item in toc:
    add_para(item, color=DARK, size=11)

doc.add_page_break()

# ================================================================
# SECTION 1: Overview
# ================================================================
add_heading_styled('1. نظرة عامة وأهداف الخطة', level=1)

add_para('تهدف هذه الخطة إلى توحيد 5 أنظمة متفرقة في 3 مكونات موحدة مع Wrapper مركزي (AppShell)، مع اتباع نهج تدريجي آمن يضمن عدم كسر النظام في أي مرحلة.', bold=True, color=DARK)

add_para('')
add_para('الأنظمة الحالية:', bold=True, color=DARK_BLUE)
items = [
    'Sidebar.tsx (579 سطر) — مستخدم في 35 صفحة مع عزل أدوار',
    'AdminSidebar.tsx (75 سطر) — مستخدم في صفحة /admin فقط',
    'Header.tsx (182 سطر) — ساعة رقمية + عنوان',
    'AdminNavbar.tsx (46 سطر) — هامبرغر + بحث + جرس + avatar',
    'FloatingBell.tsx (659 سطر) — نظام إشعارات كامل Supabase + audio + CSRF',
]
for item in items:
    add_bullet(item, color=DARK)

add_para('')
add_para('النظام المستهدف (بعد التطوير):', bold=True, color=DARK_BLUE)
target = [
    'AppShell.tsx — Wrapper مركزي يستثني صفحتي الهبوط وتسجيل الدخول',
    'UnifiedSidebar.tsx — مكون واحد يدعم 4 أدوار (ADMIN, MANAGEMENT, TEACHER, STUDENT)',
    'UnifiedHeader.tsx — مكون واحد يجمع الساعة، الهامبرغر، البحث، الجرس، avatar',
    'useNotifications.ts — Hook مستقل للإشعارات',
    'BellNotifications.tsx — مكون الجرس',
    'Footer.tsx — يبقى كما هو ويُضاف داخل AppShell',
]
for item in target:
    add_bullet(item, color=DARK)

# ================================================================
# SECTION 2: Strangler Fig Strategy
# ================================================================
add_heading_styled('2. استراتيجية Strangler Fig Pattern', level=1)

add_para('سميت بهذا الاسم نسبة لنبات التين الخانق الذي ينمو حول الشجرة المضيفة حتى يحل محلها تدريجيًا. الفكرة: نبني النظام الجديد جنبًا إلى جنب مع القديم دون تعديله، ثم نهاجر الصفحات واحدة تلو الأخرى.', bold=True, color=DARK)

add_para('')
add_para('المبادئ الأساسية:', bold=True, color=DARK_BLUE)
principles = [
    ('لا تعديل على القديم: ', 'لا نلمس Sidebar.tsx أو Header.tsx أو FloatingBell.tsx حتى مرحلة الحذف النهائي'),
    ('بناء موازي: ', 'جميع الملفات الجديدة تكون في مسار منفصل src/components/appshell/'),
    ('هجرة صفحة بصفحة: ', 'كل صفحة تتحول من النظام القديم إلى الجديد بشكل منفصل'),
    ('اختبار بعد كل هجرة: ', 'بعد تحويل كل صفحة نختبرها يدويًا قبل الانتقال للتالية'),
    ('آلية استرجاع: ', 'نحتفظ بالقديم حتى نهاية الخطة — استرجاع صفحة = تغيير سطر واحد'),
    ('صفر تعارض: ', 'لأن القديم والجديد لا يشتركان في نفس المكونات، لا يوجد تعارض'),
]
for bold_part, text in principles:
    add_bullet(text, color=DARK, bold_prefix=bold_part)

# ================================================================
# SECTION 3: Phase 1 - Preparation
# ================================================================
add_heading_styled('3. المرحلة الأولى — التحضير والتأهب (Preparation)', level=1)

add_para('قبل كتابة أي كود، نؤسس البيئة ونتأكد من فهم كل التفاصيل.', bold=True, color=DARK)

add_heading_styled('3.1  تحليل التبعيات', level=2)
deps = [
    'حصر جميع الصفحات التي تستخدم Sidebar.tsx حالياً (35 صفحة)',
    'حصر جميع الصفحات التي تستخدم Header.tsx حالياً (27 صفحة)',
    'تحديد الصفحات التي تستخدم FloatingBell.tsx وإلى أي مدى',
    'فهم تدفق uiStore: من يستدعي toggleSidebar() وأين',
    'تحديد جميع مكونات AdminNavbar التي تعتمد على props خارجية',
]
for item in deps:
    add_bullet(item, color=DARK)

add_heading_styled('3.2  إنشاء مجلدات جديدة', level=2)
add_para('ننشئ مسارين موازيين للمكونات الجديدة:', bold=True, color=DARK)
paths = [
    'src/components/appshell/ — للمكونات الجديدة (UnifiedSidebar, UnifiedHeader, AppShell)',
    'src/hooks/notifications/ — للـ useNotifications hook',
    'src/components/notifications/ — لـ BellNotifications',
]
for item in paths:
    add_bullet(item, color=DARK)

add_heading_styled('3.3  تحديث uiStore (إن لزم)', level=2)
add_para('في هذه المرحلة، قد نضيف حالات جديدة إلى uiStore (مثلاً searchOpen) لكن دون التأثير على الوظائف الحالية.', color=DARK)

# ================================================================
# SECTION 4: Phase 2 - UnifiedSidebar (Parallel)
# ================================================================
add_heading_styled('4. المرحلة الثانية — بناء UnifiedSidebar (جنبًا إلى جنب)', level=1)

add_para('هذه أهم مرحلة. نبني UnifiedSidebar.tsx في المسار الجديد دون المساس بـ Sidebar.tsx القديم.', bold=True, color=DARK)

add_heading_styled('4.1  ماذا ننقل من Sidebar.tsx', level=2)
transfer1 = [
    'مصفوفة allMenuItems كاملة (20 عنصر) مع توسيعها لدور STUDENT',
    'آلية عزل الأدوار (effectiveRoles + roles array)',
    'seenLabels Set لمنع تكرار العناصر',
    'Framer Motion: AnimatePresence, stagger, layoutId="activeDot"',
    'User card: اسم المستخدم + أول حرف + شارة المستوى + شارة الإدارة',
    'Logout: POST /api/auth/logout + unsubscribePushNotifications بالتوازي',
    'التصميم البصري بالكامل (ألوان، حواف، أيقونات)',
]
for item in transfer1:
    add_bullet(item, color=DARK)

add_heading_styled('4.2  ماذا ننقل من AdminSidebar.tsx', level=2)
transfer2 = [
    'آلية responsive: 1024px كحد فاصل (desktop ثابت / mobile overlay)',
    'إغلاق تلقائي عند تغيير المسار (useEffect على pathname)',
    'حالة الاتصال (🟢 "حالة النظام · 27 يوماً")',
    'خلفية شفافة مع backdrop blur',
]
for item in transfer2:
    add_bullet(item, color=DARK)

add_heading_styled('4.3  الإضافات الجديدة', level=2)
new_additions = [
    'توسيع roles ليشمل STUDENT في العناصر المناسبة (الرئيسية، الإشعارات، المكتبة، محرر الأكواد، المحادثة، الإعدادات)',
    'تغيير z-index من 140 إلى 40 (تحت Header)',
    'إزالة زر toggle المدمج ونقل التحكم إلى uiStore.toggleSidebar()',
]
for item in new_additions:
    add_bullet(item, color=DARK)

add_heading_styled('4.4  هيكل ملف UnifiedSidebar.tsx (تقديري: 450-500 سطر)', level=2)
create_table(
    ['المقطع', 'الوصف', 'عدد الأسطر التقديري'],
    [
        ['imports', 'React, Next.js, Framer Motion, uiStore, useResponsive', '15'],
        ['allMenuItems', '20 عنصر + دور STUDENT', '140'],
        ['SidebarItem interface', 'نوع البيانات لكل عنصر', '10'],
        ['getFilteredItems()', 'فلترة حسب الدور مع seenLabels', '25'],
        ['UnifiedSidebar component', 'الهيكل الرئيسي مع responsive', '200-250'],
        ['UserCard', 'بطاقة المستخدم + logout', '50'],
        ['OnlineStatus', 'حالة الاتصال', '15'],
    ],
    col_widths=[4, 9, 3]
)

add_para('')
add_para('ملاحظة مهمة: في هذه المرحلة، لا نستخدم UnifiedSidebar بعد. فقط نختبر أنه يصدر (export) بشكل صحيح.', bold=True, color=DARK_ORANGE)

# ================================================================
# SECTION 5: Phase 3 - useNotifications Hook
# ================================================================
add_heading_styled('5. المرحلة الثالثة — بناء useNotifications Hook', level=1)

add_para('نستخرج كل المنطق الخدمي من FloatingBell.tsx إلى Hook مستقل.', bold=True, color=DARK)

add_heading_styled('5.1  الوظائف المنقولة', level=2)
hook_funcs = [
    'الاتصال بـ Supabase realtime channel (user-{userId})',
    'تشغيل الصوت: notification.mp3 للمواعيد، alert.mp3 للإعلانات',
    'Storm detection: > 20 إشعار خلال 10 ثوانٍ → كتم الصوت + toast واحد',
    'Toast queuing: الإشعارات تتراكم والتاب hidden → تظهر عند العودة',
    'إزالة التكرار: Set داخلي مع صلاحية 60 ثانية',
    'مراقبة حالة الاتصال (connectionState) وإعادة التحميل',
    'جلب البيانات: GET /api/notifications/list?page=1&limit=10',
    'CSRF mark-as-read: POST /api/notifications/mark-read عبر csrfFetch',
]
for item in hook_funcs:
    add_bullet(item, color=DARK)

add_heading_styled('5.2  الواجهة (API) المتوقعة', level=2)

create_table(
    ['المخرج', 'النوع', 'الوصف'],
    [
        ['notifications', 'Notification[]', 'قائمة الإشعارات الحالية'],
        ['unreadCount', 'number', 'عدد الإشعارات غير المقروءة'],
        ['loading', 'boolean', 'حالة تحميل البيانات'],
        ['error', 'string | null', 'رسالة الخطأ إن وجد'],
        ['markAsRead', '(id: string) => void', 'وضع إشعار كمقروء'],
        ['markAllAsRead', '() => void', 'وضع الكل كمقروء'],
        ['isConnected', 'boolean', 'حالة الاتصال بـ Supabase'],
        ['isStormActive', 'boolean', 'هل نحن في طور العاصفة'],
    ],
    col_widths=[4, 5, 7]
)

add_para('')
add_para('ملاحظة: الملف التقديري 200-250 سطر، ويمكن اختباره بشكل منفصل.', bold=True, color=DARK_GREEN)

# ================================================================
# SECTION 6: Phase 4 - BellNotifications Component
# ================================================================
add_heading_styled('6. المرحلة الرابعة — بناء BellNotifications Component', level=1)

add_para('مكون UI بحت يستخدم useNotifications hook.', bold=True, color=DARK)

add_heading_styled('6.1  المكونات الداخلية', level=2)
bell_parts = [
    'أيقونة الجرس مع unread badge (دائرة حمراء بعدد الإشعارات)',
    'Popup منسدل يعرض حتى 8 إشعارات',
    'كل إشعار: أيقونة + title + body + timestamp نسبي + unread dot',
    'حالة التحميل: skeleton spinner',
    'حالة الخطأ: رسالة خطأ مع زر إعادة المحاولة',
    'حالة فارغة: "لا توجد إشعارات"',
    'زر "قراءة الكل" (markAllAsRead)',
    'زر "عرض الكل" (رابط إلى /notifications)',
]
for item in bell_parts:
    add_bullet(item, color=DARK)

add_heading_styled('6.2  التصميم', level=2)
design_specs = [
    'unread: خلفية rgba(13,71,161,0.03) وحدود rgba(13,71,161,0.1)',
    'read: بدون خلفية خاصة',
    'animation: framer-motion للـ popup (slide + fade)',
    'z-index: 60 (فوق الـ Header)',
]
for item in design_specs:
    add_bullet(item, color=DARK)

add_para('')
add_para('الملف التقديري: ~150 سطر', bold=True, color=DARK_GREEN)

# ================================================================
# SECTION 7: Phase 5 - UnifiedHeader
# ================================================================
add_heading_styled('7. المرحلة الخامسة — بناء UnifiedHeader', level=1)

add_para('يجمع ميزات Header.tsx + AdminNavbar.tsx + BellNotifications.', bold=True, color=DARK)

add_heading_styled('7.1  المكونات', level=2)
header_parts = [
    'يمين: Hamburger ☰ (يستدعي toggleSidebar() من uiStore)',
    'يمين: الساعة الرقمية (Orbitron, #0D47A1, 12-hour)',
    'وسط: عنوان الموقع "سحابة الأمن السيبراني" + اسم الجامعة',
    'يسار: زر بحث 🔍 (يغير route إلى /search أو يفتح modal)',
    'يسار: BellNotifications (جرس + unread badge + popup)',
    'يسار: User avatar مع قائمة منسدلة (profile, settings, logout)',
]
for item in header_parts:
    add_bullet(item, color=DARK)

add_heading_styled('7.2  التوافق مع AdminNavbar', level=2)
add_para('AdminNavbar الحالي يستقبل props (onToggleSidebar, onOpenSearch). UnifiedHeader لن يستقبل props — بل يقرأ الحالة من uiStore مباشرة.', bold=True, color=DARK)
add_para('لذلك يجب تحديث صفحة /admin (ملف واحد) لتستخدم UnifiedHeader بدلاً من AdminNavbar.', color=DARK)

add_heading_styled('7.3  هيكل الملف التقديري: 200-250 سطر', level=2)
create_table(
    ['المقطع', 'الوصف', 'الأسطر'],
    [
        ['imports', 'React, framer-motion, useResponsive, uiStore, components', '10'],
        ['Clock', 'ساعة رقمية Orbitron مع useEffect setInterval', '30'],
        ['HamburgerButton', 'زر ☰ مع hover/tap', '15'],
        ['SearchButton', 'زر بحث', '10'],
        ['UserMenu', 'avatar + قائمة منسدلة', '50'],
        ['UnifiedHeader main', 'Layout: يمين - وسط - يسار', '60-80'],
        ['Responsive variants', 'إخفاء النص على < 640px', '15'],
    ],
    col_widths=[4, 9, 3]
)

# ================================================================
# SECTION 8: Phase 6 - AppShell
# ================================================================
add_heading_styled('8. المرحلة السادسة — بناء AppShell', level=1)

add_para('الـ Wrapper المركزي الذي يجمع كل المكونات.', bold=True, color=DARK)

add_heading_styled('8.1  هيكل AppShell.tsx (تقديري: 60-80 سطر)', level=2)

# Code block simulation
code_lines = [
    'export default function AppShell({ children }: { children: ReactNode }) {',
    '  const pathname = usePathname()',
    '  const isLandingOrLogin = pathname === "/" || pathname === "/login"',
    '  if (isLandingOrLogin) return <>{children}</>',
    '',
    '  return (',
    '    <div className="flex h-screen">',
    '      <UnifiedSidebar />',
    '      <div className="flex-1 flex flex-col">',
    '        <UnifiedHeader />',
    '        <main className="flex-1 overflow-auto">',
    '          {children}',
    '        </main>',
    '      </div>',
    '    </div>',
    '  )',
    '}',
]
for line in code_lines:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(line)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = DARK

add_para('')
add_heading_styled('8.2  كيف نستخدم AppShell؟', level=2)
add_para('في root layout (src/app/layout.tsx) نلف المكونات بـ AppShell:', color=DARK)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run = p.add_run('<AppShell><Component {...pageProps} /></AppShell>')
run.font.name = 'Consolas'
run.font.size = Pt(10)
run.font.color.rgb = DARK

# ================================================================
# SECTION 9: Phase 7 - Gradual Migration
# ================================================================
add_heading_styled('9. المرحلة السابعة — الهجرة التدريجية (صفحة صفحة)', level=1)

add_para('هنا نبدأ بتحويل الصفحات من النظام القديم إلى الجديد. هذه أخطر مرحلة لأننا نلمس النظام الفعلي.', bold=True, color=DARK)

add_heading_styled('9.1  خطة الهجرة — 3 مجموعات', level=2)

create_table(
    ['المجموعة', 'الصفحات', 'المخاطرة', 'الإجراء'],
    [
        ['المجموعة 1 (أولاً)', '/admin/*', 'منخفضة', 'تحويل صفحة /admin فقط (تستخدم AdminSidebar + AdminNavbar)'],
        ['المجموعة 2 (ثانياً)', '/dashboard, /notifications\n/library, /chat', 'متوسطة', 'تحويل الصفحات الرئيسية الأكثر استخداماً'],
        ['المجموعة 3 (أخيراً)', 'باقي 30 صفحة', 'منخفضة', 'تحويل الدفعة المتبقية بعد استقرار المجموعة 2'],
    ],
    col_widths=[3, 4, 2.5, 6.5]
)

add_heading_styled('9.2  آلية تحويل صفحة واحدة', level=2)
add_numbered('نحدد الصفحة المستهدفة', 1)
add_numbered('نزيل النظام القديم من الصفحة (نحذف import Sidebar, Header, FloatingBell)', 2)
add_numbered('نضيف الصفحة داخل AppShell (أو نضمن أنها داخله)', 3)
add_numbered('نختبر الصفحة يدويًا: كل العناصر مرئية، الروابط تعمل، الأدوار صحيحة', 4)
add_numbered('إذا نجح الاختبار → ننتقل للصفحة التالية', 5)
add_numbered('إذا فشل الاختبار → نسترجع التغيير (نعيد القديم)', 6)

add_heading_styled('9.3  آلية الاسترجاع (Rollback) لصفحة واحدة', level=2)
add_para('بما أننا نحتفظ بالملفات القديمة، استرجاع صفحة = تغيير سطر واحد فقط:', bold=True, color=DARK)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run = p.add_run('// تعليق السطر الجديد وإلغاء تعليق السطر القديم')
run.font.name = 'Consolas'
run.font.size = Pt(9)
run.font.color.rgb = GRAY
run.italic = True

# ================================================================
# SECTION 10: Phase 8 - Delete Old Systems
# ================================================================
add_heading_styled('10. المرحلة الثامنة — حذف الأنظمة القديمة', level=1)

add_para('نحذف الملفات القديمة فقط بعد التأكد من أن جميع الصفحات مهاجرة وتعمل بشكل صحيح.', bold=True, color=DARK)

add_para('')
add_para('الترتيب:', bold=True, color=DARK_BLUE)
delete_order = [
    'حذف AdminSidebar.tsx — لأن وظائفه انتقلت بالكامل إلى UnifiedSidebar',
    'حذف AdminNavbar.tsx — لأن وظائفه انتقلت بالكامل إلى UnifiedHeader',
    'حذف Sidebar.tsx — بعد التأكد من أن لا صفحة تستخدمه',
    'حذف Header.tsx — بعد التأكد من أن لا صفحة تستخدمه',
    'حذف FloatingBell.tsx — بعد التأكد من أن BellNotifications + useNotifications يغطيان كل شيء',
]
for i, item in enumerate(delete_order, 1):
    add_numbered(item, i)

add_para('')
add_para('ملاحظة: لا نحذف أي ملف حتى نتأكد من أن no imports指向ه. نستخدم grep.', bold=True, color=DARK_ORANGE)

# ================================================================
# SECTION 11: Phase 9 - Full Testing & Security Audit
# ================================================================
add_heading_styled('11. المرحلة التاسعة — اختبار شامل وتدقيق أمني', level=1)

add_heading_styled('11.1  قائمة الاختبارات', level=2)

tests = [
    ('اختبار الأدوار: ', 'تسجيل الدخول بكل دور (ADMIN, MANAGEMENT, TEACHER, STUDENT) والتأكد من ظهور العناصر الصحيحة فقط'),
    ('اختبار responsive: ', 'تصغير الشاشة إلى < 640px والتأكد من أن الـ Sidebar يتحول إلى overlay'),
    ('اختبار الإشعارات: ', 'إرسال إشعار من مستخدم آخر والتأكد من وصوله عبر Supabase realtime + صوت + badge'),
    ('اختبار التسجيل الخروج: ', 'تسجيل الخروج والتأكد من إلغاء الاشتراك في الإشعارات'),
    ('اختبار الروابط: ', 'النقر على كل عنصر قائمة والتأكد من التوجيه الصحيح'),
    ('اختبار الـ CSRF: ', 'التأكد من أن mark-as-read يستخدم csrfFetch وليس fetch عادي'),
    ('اختبار Storm detection: ', 'إرسال 25 إشعار خلال 10 ثوانٍ والتأكد من كتم الصوت'),
    ('اختبار الصفحات المستثناة: ', 'التأكد من أن / و /login لا يظهر لهم Sidebar/Header'),
    ('اختبار الـ Footer: ', 'التأكد من ظهوره في جميع الصفحات الداخلية'),
]
for bold_part, text in tests:
    add_bullet(text, color=DARK, bold_prefix=bold_part)

add_heading_styled('11.2  قائمة التحقق الأمني', level=2)
security = [
    'لا يمكن لأي مستخدم رؤية عناصر قائمة لا تنتمي لدوره',
    'الـ CSRF token يمر مع كل طلب mark-as-read',
    'الـ realtime channel مقيد بـ user-{userId} فقط',
    'الـ audio لا يُشغل تلقائياً دون موافقة المستخدم',
    'الـ logout يلغي جميع الاشتراكات',
]
for item in security:
    add_bullet(item, color=DARK)

# ================================================================
# SECTION 12: Rollback Plan
# ================================================================
add_heading_styled('12. خطة الطوارئ والاسترجاع (Rollback Plan)', level=1)

add_para('في حال ظهور مشكلة حرجة بعد الهجرة، نستطيع العودة للنظام القديم فوراً.', bold=True, color=DARK)

add_heading_styled('12.1  السيناريوهات المحتملة', level=2)

scenarios = [
    ('مشكلة في الإشعارات: ', 'إذا توقف realtime أو الـ CSRF → نسترجع FloatingBell.tsx ونعيد تفعيله في الصفحات'),
    ('مشكلة في عزل الأدوار: ', 'إذا ظهر عنصر لدور غير مخول → نسترجع Sidebar.tsx ونراجع آلية effectiveRoles'),
    ('مشكلة في الـ Layout: ', 'إذا تشوهت واجهة صفحة → نسترجع تلك الصفحة فقط لمدة ساعة'),
    ('مشكلة في الـ Build: ', 'إذا فشل الـ build → git checkout -- . في الملفات المتأثرة'),
]
for bold_part, text in scenarios:
    add_bullet(text, color=DARK, bold_prefix=bold_part)

add_heading_styled('12.2  وقت الاسترجاع', level=2)
create_table(
    ['نوع المشكلة', 'وقت الاسترجاع', 'التأثير على المستخدمين'],
    [
        ['صفحة واحدة', '1-2 دقيقة', 'تأثير على صفحة واحدة فقط'],
        ['ملف كامل (مثل FloatingBell)', '5-10 دقائق', 'قد يؤثر على 35 صفحة لكن الحل فوري'],
        ['استرجاع كامل (git checkout)', 'دقيقتان', 'يعود النظام كاملاً لما قبل الهجرة'],
    ],
    col_widths=[4, 3, 6]
)

# ================================================================
# SECTION 13: Timeline
# ================================================================
add_heading_styled('13. الجدول الزمني المقدر', level=1)

create_table(
    ['المرحلة', 'المدة المقدرة', 'الاعتماديات', 'التسليم'],
    [
        ['1 — التحضير', 'نصف يوم', 'لا يوجد', 'فهم كامل + مجلدات جاهزة'],
        ['2 — UnifiedSidebar', 'يوم ونصف', 'المرحلة 1', 'ملف جاهز للاستخدام'],
        ['3 — useNotifications', 'يوم واحد', 'المرحلة 1', 'Hook جاهز'],
        ['4 — BellNotifications', 'نصف يوم', 'المرحلة 3', 'مكون جاهز'],
        ['5 — UnifiedHeader', 'يوم واحد', 'المرحلة 2 و4', 'مكون جاهز'],
        ['6 — AppShell', 'نصف يوم', 'المرحلة 5', 'Wrapper جاهز'],
        ['7 — الهجرة (3 مجموعات)', 'يوم ونصف', 'المرحلة 6', 'كل صفحة مهاجرة'],
        ['8 — حذف القديم', 'نصف يوم', 'المرحلة 7', 'لا ملفات قديمة'],
        ['9 — اختبار شامل', 'نصف يوم', 'المرحلة 8', 'تقرير اختبار'],
        ['المجموع', '7 أيام', '', 'نظام موحد كامل'],
    ],
    col_widths=[3.5, 2.5, 3, 4.5]
)

# ================================================================
# SECTION 14: Summary
# ================================================================
doc.add_page_break()
add_heading_styled('14. الملخص', level=1)

add_para('هذه الخطة تتبع نهج Strangler Fig Pattern التدريجي الآمن. النقاط الرئيسية:', bold=True, color=DARK_BLUE)

summary_points = [
    'لا نلمس الأنظمة القديمة حتى نجهز البديل بالكامل',
    'نبني النظام الجديد جنبًا إلى جنب مع القديم (Parallel Running)',
    'نهاجر صفحة بصفحة — كل صفحة تختبر قبل الانتقال للتالية',
    'آلية استرجاع فورية خلال دقائق لأي صفحة',
    'نحذف الملفات القديمة آخر شيء — وليس قبله',
    'إجمالي المدة المقدرة: 7 أيام (مع مرونة زيادة يوم للاختبارات)',
    'نسبة خطر كسر النظام: < 10% مع الالتزام بالخطة',
]

for i, point in enumerate(summary_points, 1):
    add_numbered(point, i, bold=True)

add_para('')
add_para('بعد الانتهاء من Phase 1B، سيكون لدينا:', bold=True, color=DARK_BLUE)
final_state = [
    'ملف Sidebar واحد بدلاً من 2',
    'ملف Header واحد بدلاً من 2',
    'نظام إشعارات منظم (Hook + Component) بدلاً من 659 سطر في ملف واحد',
    'Wrapper مركزي يتحكم بظهور المكونات حسب المسار',
    'صفر تكرار في الكود — نفس المكون لكل الأدوار',
]
for item in final_state:
    add_bullet(item, color=DARK)

# ---- Save ----
output_dir = 'التقارير'
os.makedirs(output_dir, exist_ok=True)
output_path = os.path.join(output_dir, 'PHASE1B_DEVELOPMENT_PLAN.docx')
doc.save(output_path)
print(f'Document saved to: {output_path}')
print(f'Size: {os.path.getsize(output_path) / 1024:.1f} KB')
