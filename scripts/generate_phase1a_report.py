from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

doc = Document()

# ---- Page setup ----
for section in doc.sections:
    section.right_to_left = True
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.right_margin = Cm(2.5)
    section.left_margin = Cm(2.5)

# ---- Style helpers ----
style = doc.styles['Normal']
style.font.name = 'Arial'
style.font.size = Pt(11)
style.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.3
rPr = style.element.get_or_add_rPr()
rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:cs="Arial"/>')
rPr.append(rFonts)

# Dark background via section
for section in doc.sections:
    sectPr = section._sectPr
    bg = parse_xml(f'<w:bg {nsdecls("w")} w:color="1E1E2E" w:themeColor="background1"/>')
    sectPr.append(bg)

# ---- Helper functions ----
def add_heading_rtl(text, level=1):
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x00, 0xE5, 0xFF)
    return h

def add_para(text, bold=False, color=None, size=None, align_rtl=True):
    p = doc.add_paragraph()
    if align_rtl:
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(text)
    if bold:
        run.bold = True
    if color:
        run.font.color.rgb = color
    if size:
        run.font.size = Pt(size)
    return p

def add_bullet(text, level=0, color=None):
    p = doc.add_paragraph(style='List Bullet')
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.right_indent = Cm(1 + level * 0.5)
    run = p.runs[0] if p.runs else p.add_run(text)
    if not p.runs or p.runs[0].text == '':
        p.clear()
        run = p.add_run(text)
    if color:
        run.font.color.rgb = color
    return p

def add_table_row(table, cells_text, bold=False, bg_color=None):
    row = table.add_row()
    for i, (cell, txt) in enumerate(zip(row.cells, cells_text)):
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run(txt)
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
        if bold:
            run.bold = True
        if bg_color:
            shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{bg_color}"/>')
            cell._tc.get_or_add_tcPr().append(shading)
    return row

def set_cell_shading(cell, color):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

# ================================================================
# DOCUMENT CONTENT
# ================================================================

# ---- Title Page ----
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('تحليل بنية واجهة المستخدم\nPhase 1A — تقرير التدقيق المعماري')
run.font.size = Pt(26)
run.bold = True
run.font.color.rgb = RGBColor(0x00, 0xE5, 0xFF)

doc.add_paragraph()
subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('CyberSecurty v11')
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0x88, 0x88, 0xAA)

doc.add_paragraph()
info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = info.add_run('تاريخ التقرير: 19 يونيو 2026\nإجمالي الملفات التي تم تحليلها: 10 ملفات\nإجمالي أسطر الكود: 1,880+ سطر')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x99, 0x99, 0xBB)

doc.add_page_break()

# ---- Table of Contents ----
add_heading_rtl('فهرس المحتويات', level=1)
toc_items = [
    '1. نظرة عامة',
    '2. تحليل نظام Sidebar',
    '   2.1  Sidebar.tsx — 579 سطر',
    '   2.2  AdminSidebar.tsx — 75 سطر',
    '   2.3  جدول مقارنة Sidebar',
    '3. تحليل نظام Header',
    '   3.1  Header.tsx — 182 سطر',
    '   3.2  AdminNavbar.tsx — 46 سطر',
    '   3.3  جدول مقارنة Header',
    '4. تحليل Footer',
    '5. تحليل FloatingBell — 659 سطر',
    '   5.1  الوظائف الرئيسية',
    '   5.2  نقاط القوة والضعف',
    '6. إدارة الحالة (State Management)',
    '7. قائمة العناصر (AllMenuItems)',
    '8. الفروقات والمشاكل الرئيسية',
    '9. خطة Phase 1B المقترحة',
    '   9.1  الهيكل الجديد',
    '   9.2  UnifiedSidebar — المواصفات',
    '   9.3  UnifiedHeader — المواصفات',
    '   9.4  FloatingBell — إعادة الهيكلة',
    '   9.5  AppShell — المواصفات',
    '   9.6  ترتيب التنفيذ',
]
for item in toc_items:
    add_para(item, color=RGBColor(0xBB, 0xBB, 0xBB), size=10)

doc.add_page_break()

# ================================================================
# SECTION 1: Overview
# ================================================================
add_heading_rtl('1. نظرة عامة', level=1)
add_para(
    'تم إجراء تدقيق شامل لبنية واجهة المستخدم في منصة CyberSecurty v11 بهدف توحيد '
    'المكونات المتكررة وتحسين قابلية الصيانة. يغطي هذا التقرير تحليل 6 مكونات رئيسية '
    'موزعة على 10 ملفات بإجمالي يزيد عن 1,880 سطر كود.'
)
add_para(
    'الهدف النهائي (Phase 1B) هو دمج 5 أنظمة متفرقة في 3 مكونات موحدة مع wrapper مركزي '
    '(AppShell) يتحكم بالهيكل العام للمنصة.'
)

add_para('قائمة الملفات التي تم تحليلها:', bold=True)
files_list = [
    'src/components/layout/Sidebar.tsx — 579 سطر',
    'src/components/layout/Header.tsx — 182 سطر',
    'src/components/layout/Footer.tsx — 54 سطر',
    'src/components/dashboard/AdminSidebar.tsx — 75 سطر',
    'src/components/dashboard/AdminNavbar.tsx — 46 سطر',
    'src/components/ui/FloatingBell.tsx — 659 سطر',
    'src/store/uiStore.ts — 17 سطر',
    'src/hooks/useResponsive.ts — 13 سطر',
    'src/app/layout.tsx — 71 سطر',
    'src/lib/theme.ts — (للإطلاع فقط)',
]
for f in files_list:
    add_bullet(f)

# ================================================================
# SECTION 2: Sidebar Analysis
# ================================================================
add_heading_rtl('2. تحليل نظام Sidebar', level=1)
add_para(
    'يوجد حاليًا نظامان منفصلان للـ Sidebar: النظام الرئيسي (Sidebar.tsx) المستخدم '
    'في 35 صفحة، ونظام خاص بلوحة الإدارة (AdminSidebar.tsx) المستخدم في صفحة وحدة.'
)

add_heading_rtl('2.1  Sidebar.tsx — 579 سطر', level=2)
add_para('آلية عزل الأدوار:', bold=True, color=RGBColor(0x00, 0xE5, 0xFF))
add_para(
    'يستخدم النظام مصفوفة roles لكل عنصر قائمة. تبدأ آلية العزل بتحديد الصلاحيات الفعلية '
    'للمستخدم بناءً على دوره ومستوى الإدارة:'
)
add_para(
    'effectiveRoles = [userRole]\n'
    'إذا كان managementLevel من المستويات (LEVEL_1 / LEVEL_2 / LEVEL_3 / LEVEL_4) '
    'وليس null أو undefined → إضافة "MANAGEMENT" إلى المصفوفة'
)
add_para(
    'يتم تصفية عناصر القائمة حسب التداخل بين effectiveRoles و roles الخاصة بكل عنصر. '
    'توجد آلية seenLabels Set تمنع تكرار العناصر — عند وجود عنصرين بنفس التسمية '
    '(مثل "سجل العمليات" حق ADMIN و MANAGEMENT)، يظهر الأول فقط.'
)

add_para('', size=6)
add_para('مكونات Sidebar.tsx الأساسية:', bold=True, color=RGBColor(0x00, 0xE5, 0xFF))
sidebar_components = [
    '20 عنصر قائمة موزعة على 4 أدوار',
    'User card: اسم المستخدم + أول حرف من اسمه + شارة المستوى + شارة الإدارة (⭐)',
    'زر تسجيل الخروج: POST /api/auth/logout بالتوازي مع unsubscribePushNotifications',
    'Framer Motion: slide جانبي مع stagger entrance (30ms بين العناصر) و layoutId="activeDot"',
    'z-index: toggle 150 → sidebar 140 → overlay 130',
    'responsive: غير موجود — لا يوجد أي تكيف مع حجم الشاشة',
    'حالة اللوحة تُدار عبر uiStore (sidebarOpen, toggleSidebar)',
]
for c in sidebar_components:
    add_bullet(c)

add_heading_rtl('2.2  AdminSidebar.tsx — 75 سطر', level=2)
add_para(
    'نظام Sidebar مخصص فقط لصفحة /admin. يتميز ببساطته النسبية مقارنة بالنظام الرئيسي.'
)
admin_sidebar_components = [
    '16 عنصر قائمة بدون عزل أدوار (كل العناصر تظهر للأدمن)',
    'responsive: دعم كامل — mobile: overlay مع backdrop (z-index 96-97)، desktop: ثابت (z-index 95)',
    'إغلاق تلقائي عند تغيير المسار (useEffect على pathname)',
    'حالة الاتصال: 🟢 "حالة النظام · 27 يوماً"',
    'لا يحتوي على User card ولا زر تسجيل خروج',
]
for c in admin_sidebar_components:
    add_bullet(c)

add_heading_rtl('2.3  جدول مقارنة Sidebar', level=2)
table = doc.add_table(rows=1, cols=3)
table.alignment = WD_TABLE_ALIGNMENT.RIGHT
table.style = 'Table Grid'

# Header row
header_cells = table.rows[0].cells
for i, txt in enumerate(['الميزة', 'Sidebar.tsx', 'AdminSidebar.tsx']):
    header_cells[i].text = ''
    p = header_cells[i].paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(txt)
    run.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x00, 0xE5, 0xFF)
    set_cell_shading(header_cells[i], '1A1A2E')

comparison_data = [
    ['عدد الأسطر', '579', '75'],
    ['عدد العناصر', '20', '16'],
    ['عزل الأدوار', 'نعم (roles array)', 'لا'],
    ['User Card', 'نعم', 'لا'],
    ['Logout', 'نعم', 'لا'],
    ['Responsive', 'لا', 'نعم'],
    ['زخرفة (Framer Motion)', 'نعم (stagger, activeDot)', 'لا'],
    ['Online Status', 'لا', 'نعم'],
    ['عدد الصفحات المستخدم فيها', '35+', '1'],
]
for row_data in comparison_data:
    row = table.add_row()
    for i, txt in enumerate(row_data):
        row.cells[i].text = ''
        p = row.cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run(txt)
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)

# ================================================================
# SECTION 3: Header Analysis
# ================================================================
add_heading_rtl('3. تحليل نظام Header', level=1)
add_para(
    'يوجد حاليًا نظامان منفصلان للـ Header: Header.tsx الرئيسي المستخدم في 27 صفحة، '
    'وAdminNavbar.tsx المستخدم في صفحة /admin فقط.'
)

add_heading_rtl('3.1  Header.tsx — 182 سطر', level=2)
header_features = [
    'ساعة رقمية بخط Orbitron ولون #00e5ff بصيغة 12 ساعة (تحديث كل 1 ثانية)',
    'عنوان الموقع: "سحابة الأمن السيبراني"',
    'اسم الجامعة: "كلية الأمن السيبراني"',
    'تصميم مختلف للموبايل عند < 768px',
    'لا يحتوي على: زر toggle sidebar, بحث, جرس إشعارات, avatar',
    'z-index: 100',
]
for c in header_features:
    add_bullet(c)

add_heading_rtl('3.2  AdminNavbar.tsx — 46 سطر', level=2)
admin_nav_features = [
    'مكون بسيط يعتمد على Props: onToggleSidebar و onOpenSearch',
    'العناصر: hamburger ☰ + بحث 🔍 + جرس 🔔 (مع unread badge) + user avatar',
    'responsive: useResponsive() عند 640px — إخفاء النص والاسم على الموبايل',
    'framer-motion: enter animation + hover/tap على كل الأزرار',
    'نفس z-index: 100',
]
for c in admin_nav_features:
    add_bullet(c)

add_heading_rtl('3.3  جدول مقارنة Header', level=2)
table = doc.add_table(rows=1, cols=3)
table.alignment = WD_TABLE_ALIGNMENT.RIGHT
table.style = 'Table Grid'
header_cells = table.rows[0].cells
for i, txt in enumerate(['الميزة', 'Header.tsx', 'AdminNavbar.tsx']):
    header_cells[i].text = ''
    p = header_cells[i].paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(txt)
    run.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x00, 0xE5, 0xFF)
    set_cell_shading(header_cells[i], '1A1A2E')

header_comparison = [
    ['عدد الأسطر', '182', '46'],
    ['ساعة رقمية', 'نعم', 'لا'],
    ['Hamburger toggle', 'لا', 'نعم'],
    ['بحث', 'لا', 'نعم'],
    ['جرس إشعارات', 'لا', 'نعم'],
    ['User Avatar', 'لا', 'نعم'],
    ['Responsive', 'نعم (768px)', 'نعم (640px)'],
    ['عدد الصفحات', '27', '1'],
    ['z-index', '100', '100'],
]
for row_data in header_comparison:
    row = table.add_row()
    for i, txt in enumerate(row_data):
        row.cells[i].text = ''
        p = row.cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run(txt)
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)

# ================================================================
# SECTION 4: Footer Analysis
# ================================================================
add_heading_rtl('4. تحليل Footer', level=2)
add_para('ملف Footer.tsx — 54 سطر')
footer_features = [
    'ثابت في الأسفل (fixed bottom)',
    'z-index: 50',
    'خلفية ضبابية (blur(25px)) مع شفافية',
    'حقوق النشر: "محمد إبراهيم الديلمي | أحمد الهيدمة"',
    'نص فرعي: "OFFICIAL CYBER SECURITY PLATFORM © 2026" بخط Orbitron',
    'لا يحتاج عزل أدوار — يظهر للجميع',
    'لا يحتاج أي تعديل جوهري — فقط يضاف داخل AppShell',
]
for c in footer_features:
    add_bullet(c)

# ================================================================
# SECTION 5: FloatingBell Analysis
# ================================================================
add_heading_rtl('5. تحليل FloatingBell — 659 سطر', level=1)
add_para(
    'نظام الإشعارات العائم — من أكثر المكونات تعقيدًا في المنصة. يقع في الزاوية السفلية '
    'اليسرى (30px, 30px) و z-index 200.'
)
add_para(
    'يختفي في الصفحات: /, /login, /onboarding',
    color=RGBColor(0x99, 0x99, 0xBB)
)

add_heading_rtl('5.1  الوظائف الرئيسية', level=2)
bell_features = [
    'Supabase realtime channel: يستمع على channel user-{userId} من نوع notification',
    'تنبيه صوتي: ملف /sounds/notification.mp3 للمواعيد و /sounds/alert.mp3 للإعلانات',
    'Storm detection: إذا تلقى > 20 إشعار خلال 10 ثوانٍ → يكتم الصوت ويعرض toast واحد فقط',
    'Toast queuing: الإشعارات تتراكم والتاب hidden → تظهر أول ما يرجع المستخدم',
    'دعم CSRF: يستخدم csrfFetch في POST /api/notifications/mark-read',
    'إزالة التكرار: Set داخلي مع صلاحية 60 ثانية',
    'مراقبة حالة الاتصال: يتفاعل مع connectionState ويعيد التحميل عند إعادة الاتصال',
    'API endpoints: GET /api/notifications/list?page=1&limit=10',
    'popup: يعرض حتى 8 عناصر، كل عنصر: أيقونة + title + body + timestamp + unread dot',
    'unread له خلفية rgba(0,229,255,0.03) وحدود rgba(0,229,255,0.1)',
]
for c in bell_features:
    add_bullet(c)

add_heading_rtl('5.2  نقاط القوة والضعف', level=2)
add_para('نقاط القوة:', bold=True, color=RGBColor(0x00, 0xE5, 0xFF))
strengths = [
    'نظام متكامل يغطي realtime, audio, UI, security, connection resilience',
    'Storm detection يمنع إزعاج المستخدم في حال الإشعارات الكثيفة',
    'Toast queuing يحافظ على الإشعارات حتى يعود المستخدم',
    'CSRF protection يضمن أمان الـ mark-as-read',
]
for c in strengths:
    add_bullet(c, color=RGBColor(0x88, 0xCC, 0x88))

add_para('نقاط الضعف:', bold=True, color=RGBColor(0xFF, 0x66, 0x66))
weaknesses = [
    '659 سطر في ملف واحد — صيانته صعبة جدًا ويحتاج تجزئة',
    'مكرر مع icon الجرس في AdminNavbar: AdminNavbar يعرض جرس navigation فقط، وFloatingBell يعرض الـ system الكامل',
    'لا يوجد استخدام لـ React Query أو SWR — fetch يدوي',
    'Supabase channel subscription متصل مباشرة وليس عبر hook قابل لإعادة الاستخدام',
    'audio يُحمل في نفس الملف مباشرة',
]
for c in weaknesses:
    add_bullet(c, color=RGBColor(0xFF, 0x88, 0x88))

# ================================================================
# SECTION 6: State Management
# ================================================================
add_heading_rtl('6. إدارة الحالة (State Management)', level=1)

add_heading_rtl('uiStore.ts (Zustand)', level=2)
ui_store = [
    'sidebarOpen: boolean (القيمة الافتراضية: false)',
    'theme: "dark" (مثبتة بشكل ثابت — لا يمكن تغييرها)',
    'toggleSidebar(): toggle الـ sidebarOpen',
    'setSidebarOpen(open): تعيين قيمة sidebarOpen يدويًا',
]
for c in ui_store:
    add_bullet(c)

add_para(
    'الـ uiStore حاليًا بسيط جدًا ولا يغطي احتياجات النظام الجديد. '
    'سيحتاج إلى إضافة searchOpen وحالات أخرى عند بناء الـ UnifiedHeader.',
    color=RGBColor(0xBB, 0xBB, 0xBB)
)

add_heading_rtl('useResponsive.ts (Custom Hook)', level=2)
add_para('مقياس الشاشة:', bold=True)
response_ranges = [
    'mobile: < 640px',
    'tablet: 640px — 1023px',
    'desktop: >= 1024px',
]
for c in response_ranges:
    add_bullet(c)

add_para(
    'ملاحظة: توجد تباين في نقاط breakpoints بين المكونات — Header يستخدم 768px، '
    'AdminNavbar يستخدم 640px، وسيحتاج النظام الجديد إلى توحيدها.',
    color=RGBColor(0xFF, 0xCC, 0x66)
)

# ================================================================
# SECTION 7: AllMenuItems
# ================================================================
add_heading_rtl('7. قائمة العناصر (AllMenuItems)', level=1)
add_para(
    'يحتوي Sidebar.tsx على 20 عنصر قائمة في مصفوفة allMenuItems (الأسطر 18-160). '
    'توزيع العناصر حسب الأدوار:'
)

add_para('توزيع الصلاحيات:', bold=True, color=RGBColor(0x00, 0xE5, 0xFF))
permissions = [
    'مشتركة لجميع الأدوار: 5 عناصر (الرئيسية، الإشعارات، المكتبة، محرر الأكواد، المحادثة، إعدادات الحساب)',
    'ADMIN فقط: 6 عناصر (سجل العمليات، استهلاك السيرفر، رادار الأخطاء، التحكم بالبوت، كنترول المنصة)',
    'MANAGEMENT فقط: 2 عناصر (المهام المنجزة، توزيع الدرجات)',
    'ADMIN + MANAGEMENT: 6 عناصر مشتركة (إدارة التوليد، ترقية المستخدمين، حسابات مفعلة، إدارة الترم، نشر تعميم)',
    'TEACHER فقط: عنصر واحد (توزيع الدرجات)',
    'STUDENT: غير موجود — لا يوجد عنصر مخصص للطالب',
]
for c in permissions:
    add_bullet(c)

add_heading_rtl('جدول جميع العناصر', level=2)
table = doc.add_table(rows=1, cols=4)
table.alignment = WD_TABLE_ALIGNMENT.RIGHT
table.style = 'Table Grid'
header_cells = table.rows[0].cells
for i, txt in enumerate(['الرمز', 'التسمية', 'المسار', 'الصلاحية']):
    header_cells[i].text = ''
    p = header_cells[i].paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(txt)
    run.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x00, 0xE5, 0xFF)
    set_cell_shading(header_cells[i], '1A1A2E')

menu_items = [
    ['🏠', 'الرئيسية', '/dashboard', 'الكل'],
    ['🔔', 'الإشعارات', '/notifications', 'الكل'],
    ['📚', 'المكتبة', '/library', 'الكل'],
    ['💻', 'محرر الأكواد', '/code-editor', 'الكل'],
    ['💬', 'المحادثة', '/chat', 'الكل'],
    ['🏗️', 'إدارة التوليد', '/admin/generation', 'ADMIN, MANAGEMENT'],
    ['📜', 'سجل العمليات', '/admin/audit-log', 'ADMIN'],
    ['📜', 'سجل العمليات', '/management/audit-log', 'MANAGEMENT'],
    ['📋', 'المهام المنجزة', '/teacher/audit-log', 'TEACHER'],
    ['⬆️', 'ترقية المستخدمين', '/admin/promotions', 'ADMIN, MANAGEMENT'],
    ['💻', 'استهلاك السيرفر', '/admin/server-usage', 'ADMIN'],
    ['💻', 'استهلاك السيرفر', '/management/server-usage', 'MANAGEMENT'],
    ['✅', 'حسابات مفعلة', '/admin/activated-accounts', 'ADMIN, MANAGEMENT'],
    ['🛡️', 'رادار الأخطاء', '/admin/security-radar', 'ADMIN'],
    ['🤖', 'التحكم بالبوت', '/admin/bot-control', 'ADMIN'],
    ['📅', 'إدارة الترم', '/admin/semester', 'ADMIN, MANAGEMENT'],
    ['🎛️', 'كنترول المنصة', '/admin/page-control', 'ADMIN'],
    ['📢', 'نشر تعميم', '/announcements/create', 'ADMIN, MANAGEMENT, TEACHER'],
    ['📝', 'توزيع الدرجات', '/teacher/grades', 'TEACHER'],
    ['⚙️', 'إعدادات الحساب', '/settings', 'الكل'],
]
for row_data in menu_items:
    row = table.add_row()
    for i, txt in enumerate(row_data):
        row.cells[i].text = ''
        p = row.cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run(txt)
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)

# ================================================================
# SECTION 8: Issues & Differences
# ================================================================
add_heading_rtl('8. الفروقات والمشاكل الرئيسية', level=1)

table = doc.add_table(rows=1, cols=3)
table.alignment = WD_TABLE_ALIGNMENT.RIGHT
table.style = 'Table Grid'
header_cells = table.rows[0].cells
for i, txt in enumerate(['#', 'المشكلة', 'التفاصيل']):
    header_cells[i].text = ''
    p = header_cells[i].paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(txt)
    run.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x00, 0xE5, 0xFF)
    set_cell_shading(header_cells[i], '1A1A2E')

issues_data = [
    ['1', 'Sidebar فيه toggle button خاص فيه', 'Sidebar.tsx يدمج زر الـ toggle (44px, cyan glow) داخل نفسه، بينما AdminSidebar يستقبل onClose prop من الخارج'],
    ['2', 'AdminNavbar غير متوافق مع Sidebar.tsx', 'AdminNavbar مبرمج لـ AdminSidebar عبر props (onToggleSidebar, onOpenSearch)، وليس مع Sidebar.tsx و uiStore'],
    ['3', 'Header لا يرتبط بالـ Sidebar', 'Header.tsx القديم لا يحتوي على hamburger, بحث, أو جرس — فقط ساعة وعنوان'],
    ['4', 'Breakpoints غير متطابقة', 'Header: 768px, AdminNavbar: 640px, Sidebar: لا يوجد'],
    ['5', 'FloatingBell مكرر الوظائف', 'جرس في AdminNavbar (navigation فقط) + FloatingBell (نظام كامل) — تداخل في الوظائف'],
    ['6', 'FloatingBell معقد جدًا (659 سطر)', 'يجمع realtime, audio, storm detection, toast queue, csrf fetch, connection monitoring, dedup في ملف واحد'],
    ['7', 'AdminSidebar يفتقر إلى User Card و Logout', 'بينما Sidebar.tsx الرئيسي يحتوي على avatar, name, role, level, management badge, و logout كامل'],
    ['8', 'لوحة ألوان مزدوجة', 'توجد لوحتان ألوان مختلفتان: tailwind.config.ts (28 لون) vs src/lib/theme.ts (21 لون) — لم يتم حلها بعد'],
    ['9', 'نقص دور STUDENT', 'لا يوجد أي عنصر مخصص لدور الطالب في allMenuItems'],
    ['10', 'uiStore محدود جدًا', 'يحتوي فقط على sidebarOpen — لا يغطي searchOpen ولا حالات أخرى مطلوبة'],
]
for row_data in issues_data:
    row = table.add_row()
    for i, txt in enumerate(row_data):
        row.cells[i].text = ''
        p = row.cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run(txt)
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
        if i == 0:
            run.font.color.rgb = RGBColor(0x00, 0xE5, 0xFF)
            run.bold = True

# ================================================================
# SECTION 9: Phase 1B Plan
# ================================================================
add_heading_rtl('9. خطة Phase 1B المقترحة', level=1)
add_para(
    'الهدف: دمج 5 أنظمة متفرقة في 3 مكونات موحدة مع Wrapper مركزي (AppShell)، '
    'مع الحفاظ على 100% من المظهر البصري الحالي (Zero Regression).',
    bold=True
)

add_heading_rtl('9.1  الهيكل الجديد', level=2)
add_para('الهيكل العام للنظام بعد التطوير:')
structure = [
    'AppShell (Wrapper)',
    '├── UnifiedSidebar (يحل محل Sidebar.tsx + AdminSidebar.tsx)',
    '├── UnifiedHeader (يحل محل Header.tsx + AdminNavbar.tsx + جرس FloatingBell)',
    '│   └── BellNotifications (مكون الجرس المستخرج من FloatingBell)',
    '├── <main>المحتوى</main>',
    '└── Footer (باقي كما هو مع إضافة داخل AppShell)',
]
for line in structure:
    add_para(line, color=RGBColor(0xBB, 0xBB, 0xBB), size=9)

add_heading_rtl('9.2  UnifiedSidebar — المواصفات', level=2)
sidebar_spec = [
    'دمج وظائف Sidebar.tsx + AdminSidebar.tsx في مكون واحد',
    'توسيع roles ليشمل STUDENT (حاليًا ADMIN/MANAGEMENT/TEACHER موجودة)',
    'إضافة responsive breakpoint عند 1024px: desktop → ثابت, mobile → overlay',
    'نقل زر toggle من داخل الـ Sidebar إلى الـ Header (hamburger icon)',
    'الحفاظ على Framer Motion: stagger entrance + activeDot + user card + logout',
    'z-index: 40 (تحت Header 50 وفوق المحتوى)',
    'إلغاء AdminSidebar.tsx بالكامل',
    'استخدام uiStore.toggleSidebar() للتحكم في حالة الفتح/الإغلاق',
]
for c in sidebar_spec:
    add_bullet(c)

add_heading_rtl('9.3  UnifiedHeader — المواصفات', level=2)
header_spec = [
    'الحفاظ على الساعة الرقمية (Orbitron, #00e5ff) من Header.tsx',
    'إضافة hamburger ☰ الذي يستدعي toggleSidebar() من uiStore',
    'إضافة زر بحث (يفتح modal أو route /search)',
    'دمج BellNotifications (جرس + unread badge + popup)',
    'إضافة user avatar مع قائمة منسدلة',
    'استخدام useResponsive() عند 640px (breakpoint واحد موحد)',
    'z-index: 50 (فوق الـ Sidebar 40)',
    'إلغاء Header.tsx و AdminNavbar.tsx بالكامل',
]
for c in header_spec:
    add_bullet(c)

add_heading_rtl('9.4  FloatingBell — إعادة الهيكلة', level=2)
add_para('تقسيم FloatingBell.tsx (659 سطر) إلى جزئين:')

add_para('useNotifications Hook (~200-250 سطر)', bold=True, color=RGBColor(0x00, 0xE5, 0xFF))
hook_items = [
    'Supabase realtime channel (user-{userId})',
    'Audio play (new-message.mp3, alert.mp3)',
    'Storm detection (> 20 إشعار في 10 ثوانٍ)',
    'Toast queuing مع إخفاء التاب',
    'Connection monitoring وإعادة التحميل',
    'Deduplication Set مع 60s expiry',
    'جلب البيانات عبر GET /api/notifications/list',
    'CSRF mark-as-read عبر POST /api/notifications/mark-read',
]
for c in hook_items:
    add_bullet(c)

add_para('BellNotifications Component (~150 سطر)', bold=True, color=RGBColor(0x00, 0xE5, 0xFF))
bell_items = [
    'أيقونة الجرس مع unread badge (عدد الإشعارات غير المقروءة)',
    'Popup يعرض حتى 8 عناصر',
    'كل عنصر: أيقونة + title + body + timestamp + unread dot',
    'حالات: تحميل, خطأ, قائمة فارغة',
    'unread تصميم: خلفية rgba(0,229,255,0.03) وحدود rgba(0,229,255,0.1)',
]
for c in bell_items:
    add_bullet(c)

add_para(
    'المجموع المقدر: ~400 سطر بدل 659 سطر — تقليل بنسبة ~40%',
    bold=True, color=RGBColor(0x88, 0xCC, 0x88)
)

add_heading_rtl('9.5  AppShell — المواصفات', level=2)
add_para('الكود المفاهيمي:')
code = (
    'AppShell ({ children }) {\n'
    '  const pathname = usePathname()\n'
    '  const isLanding = pathname === "/"\n'
    '  const isLogin = pathname === "/login"\n'
    '  if (isLanding || isLogin) return children\n\n'
    '  return (\n'
    '    <div className="flex h-screen">\n'
    '      <UnifiedSidebar />\n'
    '      <div className="flex-1 flex flex-col">\n'
    '        <UnifiedHeader />\n'
    '        <main className="flex-1 overflow-auto">{children}</main>\n'
    '        <Footer />\n'
    '      </div>\n'
    '    </div>\n'
    '  )\n'
    '}'
)
add_para(code, color=RGBColor(0xBB, 0xBB, 0xBB), size=9)

add_heading_rtl('9.6  ترتيب التنفيذ', level=2)

table = doc.add_table(rows=1, cols=3)
table.alignment = WD_TABLE_ALIGNMENT.RIGHT
table.style = 'Table Grid'
header_cells = table.rows[0].cells
for i, txt in enumerate(['الخطوة', 'الملف', 'الوصف']):
    header_cells[i].text = ''
    p = header_cells[i].paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(txt)
    run.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x00, 0xE5, 0xFF)
    set_cell_shading(header_cells[i], '1A1A2E')

exec_plan = [
    ['1', 'UnifiedSidebar.tsx', 'بناء المكون الموحد (دمج Sidebar.tsx + AdminSidebar.tsx) مع responsive 1024px'],
    ['2', 'useNotifications.ts', 'استخراج hook الإشعارات من FloatingBell.tsx'],
    ['3', 'BellNotifications.tsx', 'بناء مكون الجرس الجديد'],
    ['4', 'UnifiedHeader.tsx', 'بناء المكون الموحد (دمج Header.tsx + AdminNavbar.tsx + BellNotifications)'],
    ['5', 'AppShell.tsx', 'بناء الـ wrapper + تطبيق route groups'],
    ['6', 'تطبيق AppShell', 'تطبيق الـ AppShell على route groups — landing/login بدون، الباقي مع'],
    ['7', 'حذف القديم', 'حذف: Sidebar.tsx, AdminSidebar.tsx, Header.tsx, AdminNavbar.tsx, FloatingBell.tsx'],
    ['8', 'Footer', 'تعديل Footer وإضافته داخل AppShell'],
    ['9', 'اختبار', 'اختبار شامل لجميع الأدوار والصفحات مع التأكد من Zero Regression'],
]
for row_data in exec_plan:
    row = table.add_row()
    for i, txt in enumerate(row_data):
        row.cells[i].text = ''
        p = row.cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run(txt)
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
        if i == 0:
            run.font.color.rgb = RGBColor(0x00, 0xE5, 0xFF)
            run.bold = True

# ================================================================
# Final Notes
# ================================================================
doc.add_page_break()
add_heading_rtl('خاتمة', level=1)
add_para(
    'هذا التقرير يوثق الوضع الحالي لبنية واجهة المستخدم بالكامل ويقدم خطة واضحة '
    'للمرحلة القادمة (Phase 1B). الهدف الأساسي هو توحيد المكونات المتكررة وتقليل '
    'التعقيد مع ضمان عدم حدوث أي تراجع في المظهر البصري.'
)
add_para(
    'بعد اعتماد الخطة من الفريق، سيتم البدء في التنفيذ حسب ترتيب الخطوات المذكور. '
    'كل خطوة ستكون قابلة للاختبار بشكل منفصل لضمان الجودة قبل الانتقال للخطوة التالية.'
)

# ---- Save ----
output_dir = 'التقارير'
os.makedirs(output_dir, exist_ok=True)
output_path = os.path.join(output_dir, 'PHASE1A_UI_ARCHITECTURE_AUDIT_FULL.docx')
doc.save(output_path)
print(f'Document saved to: {output_path}')
print(f'Size: {os.path.getsize(output_path) / 1024:.1f} KB')
